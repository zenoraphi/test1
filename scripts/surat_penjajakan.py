import sys
import os
from docx import Document

def replace_in_docx(doc, data):
    """Replace placeholders in DOCX with actual data"""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            placeholder1 = f"{{{{ {key} }}}}"
            placeholder2 = f"{{{{{key}}}}}"
            
            if placeholder1 in paragraph.text or placeholder2 in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder1, str(value))
                paragraph.text = paragraph.text.replace(placeholder2, str(value))
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    placeholder1 = f"{{{{ {key} }}}}"
                    placeholder2 = f"{{{{{key}}}}}"
                    
                    if placeholder1 in cell.text or placeholder2 in cell.text:
                        cell.text = cell.text.replace(placeholder1, str(value))
                        cell.text = cell.text.replace(placeholder2, str(value))
    
    return doc

def add_page_break(doc):
    """Add page break"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_break()

def merge_documents(main_doc, append_doc):
    """Merge two documents"""
    add_page_break(main_doc)
    
    for element in append_doc.element.body:
        main_doc.element.body.append(element)
    
    return main_doc

def convert_to_pdf(docx_path, pdf_path):
    """Convert DOCX to PDF using LibreOffice"""
    try:
        import subprocess
        import shutil
        
        # Cek apakah libreoffice/soffice ada di PATH
        libreoffice_cmd = shutil.which('libreoffice') or shutil.which('soffice.exe')
        
        if not libreoffice_cmd:
            print("[ERROR] LibreOffice tidak ditemukan di PATH!")
            print("[INFO] Install LibreOffice dan tambahkan ke PATH")
            return False
        
        # Command untuk convert
        cmd = [
            libreoffice_cmd,
            '--headless',
            '--convert-to',
            'pdf',
            '--outdir',
            os.path.dirname(pdf_path),
            docx_path
        ]
        
        print(f"[INFO] Running: {' '.join(cmd)}")
        
        # Jalankan command
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0:
            # LibreOffice bikin file di folder yang sama
            generated_pdf = docx_path.replace('.docx', '.pdf')
            if generated_pdf != pdf_path and os.path.exists(generated_pdf):
                shutil.move(generated_pdf, pdf_path)
            print(f"[SUCCESS] PDF created: {pdf_path}")
            return True
        else:
            print(f"[ERROR] LibreOffice failed: {result.stderr}")
            return False
            
    except Exception as e:
        print(f"[ERROR] PDF conversion error: {str(e)}")
        return False

def main():
    if len(sys.argv) != 14:
        print("[ERROR] Incorrect number of arguments")
        print("Usage: python surat_penjajakan.py <template_path> <lampiran_path> <output_path> <tanggal_surat> <nomor_surat> <nama_dudi> <alamat_jalan> <alamat_kecamatan> <lama_pkl> <tanggal_mulai> <tanggal_selesai> <tingkat> <jurusan>")
        sys.exit(1)
    
    # Parse arguments
    template_path = sys.argv[1]
    lampiran_path = sys.argv[2]
    output_path = sys.argv[3]
    
    data = {
        'tanggal_surat': sys.argv[4],
        'nomor_surat': sys.argv[5],
        'nama_dudi': sys.argv[6],
        'alamat_jalan': sys.argv[7],
        'alamat_kecamatan': sys.argv[8],
        'lama_pkl': sys.argv[9],
        'tanggal_mulai': sys.argv[10],
        'tanggal_selesai': sys.argv[11],
        'tingkat': sys.argv[12],
        'jurusan': sys.argv[13],
        'pembekalan': '(akan diisi)',
        'pengujian': '(akan diisi)'
    }
    
    print(f"[INFO] Starting script with data: {data}")
    
    # Check templates
    if not os.path.exists(template_path):
        print(f"[ERROR] Template file not found: {template_path}")
        sys.exit(1)
    
    if not os.path.exists(lampiran_path):
        print(f"[ERROR] Lampiran file not found: {lampiran_path}")
        sys.exit(1)
    
    try:
        # Load & process documents
        print(f"[INFO] Loading template: {template_path}")
        main_doc = Document(template_path)
        main_doc = replace_in_docx(main_doc, data)
        
        print(f"[INFO] Loading lampiran: {lampiran_path}")
        lampiran_doc = Document(lampiran_path)
        
        print("[INFO] Merging documents...")
        merged_doc = merge_documents(main_doc, lampiran_doc)
        
        # Save DOCX
        docx_output = output_path + '.docx'
        merged_doc.save(docx_output)
        print(f"[OK] DOCX created: {docx_output}")
        
        # Convert to PDF
        pdf_output = output_path + '.pdf'
        print(f"[INFO] Converting to PDF: {pdf_output}")
        
        if convert_to_pdf(docx_output, pdf_output):
            print(f"[OK] PDF created: {pdf_output}")
        else:
            print("[ERROR] PDF conversion failed!")
            sys.exit(1)
        
        print("[SUCCESS] Document processed successfully")
        sys.exit(0)
        
    except Exception as e:
        print(f"[ERROR] Failed to process document: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()